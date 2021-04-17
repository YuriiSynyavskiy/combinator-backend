import os
from docx import Document
import nltk
import re
nltk.download('averaged_perceptron_tagger')


class AnalyzeText():
    def __init__(self, file, filters):
        file.save(file.filename)
        self.path = file.filename
        self.filters = [item for item in filters.values()]

    def open_document(self):
        return Document(self.path)

    def calculate_paragraphs(self, document):
        return len(document.paragraphs)

    def apply_filters(self, document):
        i = 0
        matched_data = []
        while i < self.calculate_paragraphs(document):
            this_paragraph_words = [word for word in document.paragraphs[i].text.split(
            ) if word and len(word) > 1]

            processed_data = nltk.pos_tag(this_paragraph_words)
            k = 0
            paragraph_len = len(processed_data)
            while(k < paragraph_len):
                j = k
                for word_filter in self.filters:
                    if j < paragraph_len and self.verify_word_with_filters(processed_data[j], word_filter):
                        j += 1
                    else:
                        break
                if j-len(self.filters) == k:
                    matching = processed_data[k:j]
                    words = ' '.join([re.sub(r'[^\w]', '', word[0]) for word in matching])
                    words_types = '  '.join([word[1] for word in matching])
                    matched_data.append([words, words_types])
                k += 1
            i+=1
        return matched_data

    def analyze_file(self):
        document = self.open_document()
        result = self.apply_filters(document)
        os.remove(self.path)
        return result

    def verify_word_with_filters(self, word, filters):
        if word[1] in filters.split(","):
            return True